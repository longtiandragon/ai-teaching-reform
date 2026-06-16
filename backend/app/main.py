import asyncio
import json
from pathlib import Path
from contextlib import asynccontextmanager

from fastapi import FastAPI, File, Form, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse

from backend.app.config import get_settings
from backend.app.data.courses import NONGBO_COURSE_ID, SPRINGBOOT_COURSE_ID, get_courses, get_lesson
from backend.app.models import (
    ChatRequest,
    ChatResponse,
    AIChatTaskRequest,
    AICheckRequest,
    AICheckResponse,
    AgentSkillRunRequest,
    AgentSkillRunResponse,
    CourseLineSummary,
    GuidedMessageRequest,
    GuidedSessionResponse,
    GuidedStartRequest,
    HealthResponse,
    IngestResponse,
    LearningMapResponse,
    LearningRecord,
    LearningTaskDetail,
    PracticeFeedback,
    PracticeSubmission,
    SelfCheckSubmission,
    SessionBootstrapResponse,
    SessionLoginRequest,
    SessionLoginResponse,
    UserInfo,
)
from backend.app.services.auth import create_access_token, verify_password
from backend.app.services.analytics import teacher_analytics
from backend.app.services.agent_tools import list_skills, run_skill
from backend.app.services.agent_flow import NoCitationError, run_agent_once, stream_agent
from backend.app.services.guided_agent import (
    continue_guided_session,
    finish_guided_continue,
    finish_guided_start,
    prepare_guided_continue,
    prepare_guided_start,
    start_guided_session,
    stream_teaching_step,
)
from backend.app.services.ai_check import check_learning_task
from backend.app.services.course_runtime import (
    get_course_line,
    get_learning_map,
    get_task_detail,
    list_course_lines,
    sync_course_runtime_seed,
)
from backend.app.services.database import (
    db_exists,
    get_user_by_login,
    get_user,
    init_db,
    list_classes,
    list_users,
    record_interaction,
    record_learning_event,
    record_task_submission,
    student_records,
    list_deleted_question_ids,
    list_teacher_questions,
    get_teacher_question,
    soft_delete_teacher_question,
    sync_published_questions_seed,
    upsert_teacher_question,
)
from backend.app.services.llm import LLMError, deepseek_client
from backend.app.services.practice import evaluate_practice
from backend.app.services.rag import rag_service
from backend.app.services import kb_manager, ai_config


@asynccontextmanager
async def lifespan(_: FastAPI):
    init_db()
    sync_course_runtime_seed()
    sync_published_questions_seed(
        [question.model_dump() for question in questions_by_course(SPRINGBOOT_COURSE_ID)]
        + [question.model_dump() for question in questions_by_course(NONGBO_COURSE_ID)]
    )
    rag_service.ingest_seed()
    # 同步种子文件到知识库文件表
    kb_manager.sync_seed_files_to_kb(rag_service._chunks)
    yield


app = FastAPI(title="Web Training Platform", version="0.1.0", lifespan=lifespan)
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173", "http://127.0.0.1:5173"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.get("/api/health", response_model=HealthResponse)
async def health() -> HealthResponse:
    settings = get_settings()
    return HealthResponse(
        status="ok",
        database=db_exists(),
        rag_backend=rag_service.backend_name,
        rag_config=rag_service.status(),
        deepseek_configured=deepseek_client.configured,
        deepseek_live=settings.deepseek_live,
        model=settings.deepseek_model,
    )


@app.get("/api/courses")
async def courses() -> dict:
    return {"courses": get_courses()}


@app.get("/api/courses/{course_id}/lessons/{lesson_id}")
async def lesson(course_id: str, lesson_id: str) -> dict:
    return {"lesson": get_lesson(course_id, lesson_id)}


@app.get("/api/course-lines", response_model=list[CourseLineSummary])
async def course_lines() -> list[CourseLineSummary]:
    return await asyncio.to_thread(list_course_lines)


@app.get("/api/course-lines/{course_line_id}", response_model=CourseLineSummary)
async def course_line(course_line_id: str) -> CourseLineSummary:
    return await asyncio.to_thread(get_course_line, course_line_id)


@app.get("/api/course-lines/{course_line_id}/learning-map", response_model=LearningMapResponse)
async def learning_map(course_line_id: str, student_id: str | None = None) -> LearningMapResponse:
    return await asyncio.to_thread(get_learning_map, course_line_id, student_id)


@app.get("/api/tasks/{task_id}", response_model=LearningTaskDetail)
async def task_detail(task_id: str) -> LearningTaskDetail:
    return await asyncio.to_thread(get_task_detail, task_id)


@app.get("/api/session/bootstrap", response_model=SessionBootstrapResponse)
async def session_bootstrap() -> SessionBootstrapResponse:
    classes = [dict(row) for row in await asyncio.to_thread(list_classes)]
    class_names = {item["id"]: item["name"] for item in classes}
    users = []
    for row in await asyncio.to_thread(list_users):
        item = dict(row)
        item["class_name"] = class_names.get(item.get("class_id"))
        users.append(item)
    return SessionBootstrapResponse(classes=classes, users=users)


@app.post("/api/session/login", response_model=SessionLoginResponse)
async def session_login(request: SessionLoginRequest) -> SessionLoginResponse:
    row = await asyncio.to_thread(get_user_by_login, request.account)
    if not row or not verify_password(request.password, row["password_hash"]):
        raise HTTPException(status_code=401, detail="账号或密码错误")
    item = dict(row)
    classes = {row["id"]: row["name"] for row in await asyncio.to_thread(list_classes)}
    item["class_name"] = classes.get(item.get("class_id"))
    user = UserInfo(**item)
    return SessionLoginResponse(user=user, access_token=create_access_token(item))


@app.post("/api/chat", response_model=ChatResponse)
async def chat(request: ChatRequest) -> ChatResponse:
    try:
        result = await run_agent_once(request)
    except NoCitationError as exc:
        raise HTTPException(status_code=422, detail=str(exc)) from exc
    except LLMError as exc:
        raise HTTPException(status_code=exc.status_code, detail=str(exc)) from exc
    return ChatResponse(
        answer=result.get("answer", ""),
        citations=result.get("citations", []),
        suggestions=[result.get("next_action", "")] if result.get("next_action") else [],
        fallback=result.get("fallback", False),
    )


@app.post("/api/ai/chat", response_model=ChatResponse)
async def ai_chat(request: AIChatTaskRequest) -> ChatResponse:
    chat_request = ChatRequest(
        course_id=request.courseLineId,
        lesson_id=request.taskId,
        question=request.question,
        mode="student",
        student_id=request.studentId,
    )
    return await chat(chat_request)


@app.post("/api/ai/hint", response_model=ChatResponse)
async def ai_hint(request: AIChatTaskRequest) -> ChatResponse:
    request.question = f"请只围绕当前任务给一到两步提示，不要直接给完整答案。问题：{request.question}"
    return await ai_chat(request)


@app.post("/api/ai/reflect", response_model=ChatResponse)
async def ai_reflect(request: AIChatTaskRequest) -> ChatResponse:
    request.question = f"请帮助学生基于当前任务生成学习复盘，不要编造资料。学生内容：{request.question}"
    return await ai_chat(request)


@app.post("/api/ai/check", response_model=AICheckResponse)
async def ai_check(request: AICheckRequest) -> AICheckResponse:
    try:
        return await check_learning_task(request)
    except NoCitationError as exc:
        raise HTTPException(status_code=422, detail=str(exc)) from exc
    except LLMError as exc:
        raise HTTPException(status_code=exc.status_code, detail=str(exc)) from exc


# ---------- AI 配置 API ----------

@app.get("/api/ai/config")
async def get_ai_config_endpoint() -> dict:
    return ai_config.get_ai_config()


@app.put("/api/ai/config")
async def update_ai_config_endpoint(payload: dict) -> dict:
    return ai_config.update_ai_config(
        api_key=payload.get("apiKey"),
        base_url=payload.get("baseUrl"),
        model=payload.get("model"),
        live=payload.get("live"),
        timeout=payload.get("timeout"),
    )


@app.get("/api/ai/presets")
async def get_ai_presets() -> dict:
    return {"presets": ai_config.get_presets()}


@app.post("/api/ai/test")
async def test_ai_connection() -> dict:
    try:
        from openai import AsyncOpenAI
        settings = get_settings()
        client = AsyncOpenAI(api_key=settings.deepseek_api_key, base_url=settings.deepseek_base_url)
        response = await asyncio.wait_for(
            client.chat.completions.create(
                model=settings.deepseek_model,
                messages=[{"role": "user", "content": "say hello"}],
                max_tokens=10,
            ),
            timeout=10,
        )
        text = response.choices[0].message.content or ""
        return {"ok": True, "message": "连接正常", "response": text[:100]}
    except Exception as exc:
        return {"ok": False, "message": f"连接失败: {str(exc)}"}


@app.post("/api/chat/stream")
async def chat_stream(request: ChatRequest) -> StreamingResponse:
    async def event_stream():
        try:
            async for event, data in stream_agent(request):
                yield _sse(event, data)
        except NoCitationError as exc:
            yield _sse("error", {"message": str(exc), "status": 422})
        except LLMError as exc:
            yield _sse("error", {"message": str(exc), "status": exc.status_code})
        except Exception as exc:
            yield _sse("error", {"message": "AI 调用失败。", "detail": str(exc), "status": 500})

    return StreamingResponse(
        event_stream(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        },
    )


@app.post("/api/practice/submit", response_model=PracticeFeedback)
async def practice_submit(request: PracticeSubmission) -> PracticeFeedback:
    try:
        feedback = await evaluate_practice(request.course_id, request.lesson_id, request.code, request.notes)
    except NoCitationError as exc:
        raise HTTPException(status_code=422, detail=str(exc)) from exc
    except LLMError as exc:
        raise HTTPException(status_code=exc.status_code, detail=str(exc)) from exc
    await asyncio.to_thread(
        record_interaction,
        kind="practice",
        course_id=request.course_id,
        lesson_id=request.lesson_id,
        question=request.notes or "practice submission",
        answer=feedback.ai_comment,
        score=feedback.score,
        student_id=request.student_id,
    )
    if request.student_id:
        await asyncio.to_thread(
            record_learning_event,
            student_id=request.student_id,
            course_id=request.course_id,
            lesson_id=request.lesson_id,
            kind="ai_practice",
            score=feedback.score,
            code=request.code,
            notes=request.notes,
            feedback=feedback.ai_comment,
        )
    return feedback


@app.post("/api/learning/self-check")
async def learning_self_check(request: SelfCheckSubmission) -> dict:
    await asyncio.to_thread(
        record_learning_event,
        student_id=request.student_id,
        course_id=request.course_id,
        lesson_id=request.lesson_id,
        kind="self_check",
        score=request.score,
        correct=request.correct,
        total=request.total,
        answers=request.answers,
    )
    return {"ok": True}


@app.get("/api/students/{student_id}/records")
async def student_learning_records(student_id: str, course_id: str | None = None) -> dict[str, list[LearningRecord]]:
    rows = await asyncio.to_thread(student_records, student_id, course_id)
    return {"records": [_learning_record_from_row(row) for row in rows]}


@app.post("/api/kb/ingest", response_model=IngestResponse)
async def ingest() -> IngestResponse:
    chunks = await asyncio.to_thread(rag_service.ingest_seed)
    return IngestResponse(chunks=chunks, backend=rag_service.backend_name, message="真实课程资料知识库已就绪")


@app.post("/api/kb/upload", response_model=IngestResponse)
async def kb_upload(
    file: UploadFile = File(...),
    courseLineId: str | None = Form(default=None),
    taskId: str | None = Form(default=None),
) -> IngestResponse:
    settings = get_settings()
    safe_name = Path(file.filename or "upload.md").name
    suffix = Path(safe_name).suffix.lower()
    content_bytes = await file.read()
    if len(content_bytes) > settings.upload_max_bytes:
        raise HTTPException(status_code=413, detail=f"file too large, max {settings.upload_max_bytes} bytes")
    target = settings.upload_dir / safe_name
    await asyncio.to_thread(target.write_bytes, content_bytes)

    text_suffixes = {
        ".md", ".markdown", ".txt", ".text", ".log",
        ".java", ".py", ".js", ".ts", ".vue", ".jsx", ".tsx",
        ".xml", ".yml", ".yaml", ".json", ".sql", ".html", ".htm",
        ".css", ".scss", ".less", ".sh", ".bat", ".properties",
        ".gitignore", ".env", ".toml", ".ini", ".cfg",
        ".c", ".cpp", ".h", ".hpp", ".go", ".rs", ".rb", ".php",
    }

    try:
        if suffix == ".pdf":
            from pypdf import PdfReader
            content = await asyncio.to_thread(_extract_pdf_text, target)
        elif suffix == ".docx":
            content = await asyncio.to_thread(_extract_docx_text, target)
        elif suffix in (".xlsx", ".xls"):
            content = await asyncio.to_thread(_extract_excel_text, target)
        elif suffix == ".csv":
            content = await asyncio.to_thread(_extract_csv_text, target)
        elif suffix in text_suffixes:
            content = content_bytes.decode("utf-8", errors="ignore")
        else:
            content = content_bytes.decode("utf-8", errors="ignore")
            if not content.strip():
                raise HTTPException(status_code=400, detail=f"不支持的文件类型: {suffix}")
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(status_code=400, detail=f"文件解析失败: {str(exc)}") from exc

    file_id = await asyncio.to_thread(
        kb_manager.register_file,
        safe_name,
        suffix,
        0,
        courseLineId,
    )
    task_ids = [taskId] if taskId else []
    chunks = await asyncio.to_thread(
        rag_service.add_document,
        target.name,
        content,
        courseLineId,
        file_id=file_id,
        task_ids=task_ids,
    )
    await asyncio.to_thread(kb_manager.update_file_chunk_count, file_id, chunks)
    if taskId:
        await asyncio.to_thread(kb_manager.associate_task, file_id, taskId)
    return IngestResponse(chunks=chunks, backend=rag_service.backend_name, message="uploaded document indexed", fileId=file_id)


# ---------- 题库 API ----------

from backend.app.data.questions import SEED_QUESTIONS, questions_by_lesson, questions_by_course
from backend.app.models import Question, QuestionCreateRequest, QuestionListResponse

def _all_questions(course_id: str) -> list[Question]:
    """合并种子题 + 教师添加的题"""
    teacher = [_question_from_teacher_row(row) for row in list_teacher_questions(course_id)]
    teacher_by_id = {question.id: question for question in teacher}
    deleted_ids = list_deleted_question_ids(course_id)
    merged: list[Question] = []
    for question in questions_by_course(course_id):
        if question.id in deleted_ids:
            continue
        merged.append(teacher_by_id.pop(question.id, question))
    merged.extend(teacher_by_id.values())
    return merged


def _question_lookup() -> dict[str, Question]:
    questions = _all_questions(SPRINGBOOT_COURSE_ID) + _all_questions(NONGBO_COURSE_ID)
    return {q.id: q for q in questions}


@app.get("/api/questions", response_model=QuestionListResponse)
async def get_questions(course_id: str, lesson_id: str | None = None) -> QuestionListResponse:
    all_qs = _all_questions(course_id)
    if lesson_id:
        all_qs = [q for q in all_qs if q.lesson_id == lesson_id]
    return QuestionListResponse(questions=all_qs, total=len(all_qs))


@app.post("/api/questions", response_model=Question)
async def create_question(request: QuestionCreateRequest) -> Question:
    import uuid
    qid = f"teacher-{uuid.uuid4().hex[:8]}"
    q = Question(id=qid, **request.model_dump())
    await asyncio.to_thread(upsert_teacher_question, q.model_dump())
    return q


@app.put("/api/questions/{question_id}", response_model=Question)
async def update_question(question_id: str, request: QuestionCreateRequest) -> Question:
    if question_id not in _question_lookup():
        raise HTTPException(status_code=404, detail="question not found")
    q = Question(id=question_id, **request.model_dump())
    await asyncio.to_thread(upsert_teacher_question, q.model_dump())
    return q


@app.delete("/api/questions/{question_id}")
async def delete_question(question_id: str) -> dict:
    if await asyncio.to_thread(get_all_published_for_question, question_id):
        raise HTTPException(status_code=409, detail="published questions must be unpublished before deletion")
    existing = _question_lookup().get(question_id)
    if not existing:
        raise HTTPException(status_code=404, detail="question not found")
    if not await asyncio.to_thread(get_teacher_question, question_id):
        await asyncio.to_thread(upsert_teacher_question, existing.model_dump())
    ok = await asyncio.to_thread(soft_delete_teacher_question, question_id)
    if not ok:
        raise HTTPException(status_code=404, detail="question not found")
    return {"deleted": question_id}


@app.post("/api/questions/{question_id}/generate-explanation", response_model=Question)
async def generate_question_explanation(question_id: str) -> Question:
    question = _question_lookup().get(question_id)
    if not question:
        raise HTTPException(status_code=404, detail="question not found")
    prompt = (
        "请根据题目、标准答案和项目上下文生成一段教师可直接使用的解析。"
        "要求：说明为什么答案正确，指出相关 SpringBoot/农宝项目代码或数据库依据，不要编造不存在的类名。\n"
        f"题型：{question.type}\n题干：{question.stem}\n答案：{question.answer}"
    )
    try:
        result = await run_agent_once(ChatRequest(
            course_id=question.course_id,
            lesson_id=question.lesson_id,
            question=prompt,
            mode="teacher",
        ))
        explanation = result.get("answer", "").strip()
    except (NoCitationError, LLMError) as exc:
        raise HTTPException(status_code=getattr(exc, "status_code", 422), detail=str(exc)) from exc
    updated = question.model_copy(update={"explanation": explanation})
    await asyncio.to_thread(upsert_teacher_question, updated.model_dump())
    return updated


@app.post("/api/guided/start", response_model=GuidedSessionResponse)
async def guided_start(request: GuidedStartRequest) -> GuidedSessionResponse:
    task_context, prompt = await _guided_start_context(request)
    try:
        result = await start_guided_session(
            task_id=request.taskId,
            course_line_id=request.courseLineId,
            student_id=request.studentId,
            student_input=prompt,
            task_context=task_context,
            code_draft=request.codeDraft or "",
        )
    except NoCitationError as exc:
        raise HTTPException(status_code=422, detail=str(exc)) from exc
    except LLMError as exc:
        raise HTTPException(status_code=exc.status_code, detail=str(exc)) from exc
    return GuidedSessionResponse(**result)


@app.post("/api/ai/local-check-record")
async def local_check_record(payload: dict) -> dict:
    required = ["courseLineId", "moduleId", "taskId", "studentId", "studentInput", "result"]
    if any(not payload.get(key) for key in required):
        raise HTTPException(status_code=400, detail="missing record payload")

    task = get_task_detail(str(payload["taskId"]))
    if task.course_line_id != payload["courseLineId"]:
        raise HTTPException(status_code=404, detail="task not found in course line")

    result = payload.get("result") or {}
    score = _safe_int(result.get("score"), 0)
    passed = bool(result.get("passed"))
    level = str(result.get("level") or ("passed" if passed else "needs_revision"))
    reply = str(result.get("reply") or "")
    rubric_scores = result.get("rubricScores") if isinstance(result.get("rubricScores"), list) else []
    evidence = result.get("evidence") if isinstance(result.get("evidence"), list) else []
    student_input = str(payload["studentInput"])

    await asyncio.to_thread(
        record_task_submission,
        student_id=str(payload["studentId"]),
        course_line_id=str(payload["courseLineId"]),
        module_id=str(payload["moduleId"]),
        task_id=str(payload["taskId"]),
        artifact_type=payload.get("artifactType"),
        student_input=student_input,
        score=score,
        passed=passed,
        level=level,
        reply=reply,
        rubric_scores=rubric_scores,
        evidence=evidence,
        next_task_id=None,
    )
    await asyncio.to_thread(
        record_interaction,
        kind="local_check",
        course_id=str(payload["courseLineId"]),
        lesson_id=str(payload["taskId"]),
        question=student_input,
        answer=reply,
        score=score,
        student_id=str(payload["studentId"]),
    )
    await asyncio.to_thread(
        record_learning_event,
        student_id=str(payload["studentId"]),
        course_id=str(payload["courseLineId"]),
        lesson_id=str(payload["taskId"]),
        kind="ai_check",
        score=score,
        code=student_input,
        feedback=reply,
        answers=rubric_scores,
    )

    # 自动收集错题：得分 < 70 或未通过时收录到错题本
    if score < 70 or not passed:
        # 提取知识点
        kp_list = []
        for rs in rubric_scores:
            if isinstance(rs, dict) and rs.get("name"):
                kp_list.append(rs["name"])
        knowledge_points = ", ".join(kp_list[:5])

        # 提取正确答案（从 rubric 的 answer 字段）
        correct_answer = ""
        task_detail = await asyncio.to_thread(get_task_detail, str(payload["taskId"]))
        if task_detail and hasattr(task_detail, "instruction"):
            correct_answer = task_detail.instruction[:200]

        await asyncio.to_thread(
            add_error_entry,
            student_id=str(payload["studentId"]),
            source_type="ai_check",
            source_id=f"check-{payload['taskId']}",
            course_line_id=str(payload["courseLineId"]),
            task_id=str(payload["taskId"]),
            question=student_input[:500],
            student_answer=student_input[:500],
            correct_answer=correct_answer,
            error_analysis=reply[:300] if reply else "",
            knowledge_points=knowledge_points,
            difficulty="medium",
        )

    return {"ok": True}


@app.post("/api/guided/start/stream")
async def guided_start_stream(request: GuidedStartRequest) -> StreamingResponse:
    async def event_stream():
        try:
            yield _sse("status", {"stage": "context", "message": "正在读取任务与题目..."})
            task_context, prompt = await _guided_start_context(request)
            yield _sse("status", {"stage": "planning", "message": "正在规划引导步骤..."})
            state = await prepare_guided_start(
                task_id=request.taskId,
                course_line_id=request.courseLineId,
                student_id=request.studentId,
                student_input=prompt,
                task_context=task_context,
                code_draft=request.codeDraft or "",
            )
            yield _sse("metadata", _stream_metadata(state))
            yield _sse("status", {"stage": "generating", "message": "正在生成本步引导..."})
            async for delta in stream_teaching_step(state):
                yield _sse("delta", {"text": delta})
            result = finish_guided_start(state)
            yield _sse("done", result)
        except NoCitationError as exc:
            yield _sse("error", {"message": str(exc), "status": 422})
        except LLMError as exc:
            yield _sse("error", {"message": str(exc), "status": exc.status_code})
        except Exception as exc:
            yield _sse("error", {"message": "智能体启动失败。", "detail": str(exc), "status": 500})

    return _sse_response(event_stream())


async def _guided_start_context(request: GuidedStartRequest) -> tuple[dict, str]:
    task = await asyncio.to_thread(get_task_detail, request.taskId)
    if task.course_line_id != request.courseLineId:
        raise HTTPException(status_code=404, detail="task not found in course line")
    question = None
    if request.questionId:
        question = next((item for item in _all_questions(request.courseLineId) if item.id == request.questionId), None)
        if not question:
            raise HTTPException(status_code=404, detail="question not found")
        published = await asyncio.to_thread(get_published_questions, request.taskId)
        if not any(row["question_id"] == request.questionId for row in published):
            raise HTTPException(status_code=404, detail="question is not published to this task")
    linked_files = await asyncio.to_thread(kb_manager.get_task_files, request.taskId)
    task_context = task.model_dump()
    task_context["question"] = question.model_dump() if question else None
    task_context["linkedFiles"] = linked_files
    prompt = request.studentInput
    if question and question.stem not in prompt:
        prompt = f"题目：{question.stem}\n学生请求：{prompt}"
    return task_context, prompt


@app.post("/api/guided/message", response_model=GuidedSessionResponse)
async def guided_message(request: GuidedMessageRequest) -> GuidedSessionResponse:
    try:
        result = await continue_guided_session(
            request.sessionId,
            request.message,
            code_draft=request.codeDraft or "",
        )
    except ValueError as exc:
        raise HTTPException(status_code=404, detail=str(exc)) from exc
    except NoCitationError as exc:
        raise HTTPException(status_code=422, detail=str(exc)) from exc
    except LLMError as exc:
        raise HTTPException(status_code=exc.status_code, detail=str(exc)) from exc
    return GuidedSessionResponse(**result)


@app.post("/api/guided/message/stream")
async def guided_message_stream(request: GuidedMessageRequest) -> StreamingResponse:
    async def event_stream():
        try:
            yield _sse("status", {"stage": "checking", "message": "正在判断当前输入是否可以推进..."})
            state, status = await prepare_guided_continue(
                request.sessionId,
                request.message,
                code_draft=request.codeDraft or "",
            )
            yield _sse("metadata", _stream_metadata(state))
            if state.get("held"):
                result = finish_guided_continue(state, status)
                yield _sse("done", result)
                return
            if status == "completed":
                result = finish_guided_continue(state, status)
                yield _sse("done", result)
                return
            yield _sse("status", {"stage": "generating", "message": "正在生成下一步引导..."})
            async for delta in stream_teaching_step(state):
                yield _sse("delta", {"text": delta})
            result = finish_guided_continue(state, status)
            yield _sse("done", result)
        except ValueError as exc:
            yield _sse("error", {"message": str(exc), "status": 404})
        except NoCitationError as exc:
            yield _sse("error", {"message": str(exc), "status": 422})
        except LLMError as exc:
            yield _sse("error", {"message": str(exc), "status": exc.status_code})
        except Exception as exc:
            yield _sse("error", {"message": "智能体回复失败。", "detail": str(exc), "status": 500})

    return _sse_response(event_stream())


@app.get("/api/agent/skills")
async def agent_skills() -> dict:
    return {"skills": list_skills()}


@app.post("/api/agent/run", response_model=AgentSkillRunResponse)
async def agent_run(request: AgentSkillRunRequest) -> AgentSkillRunResponse:
    try:
        return await run_skill(request.name, request.arguments)
    except KeyError as exc:
        raise HTTPException(status_code=404, detail="skill not found") from exc
    except NoCitationError as exc:
        raise HTTPException(status_code=422, detail=str(exc)) from exc
    except LLMError as exc:
        raise HTTPException(status_code=exc.status_code, detail=str(exc)) from exc


@app.get("/api/teacher/analytics")
async def analytics(student_id: str | None = None) -> dict:
    return await asyncio.to_thread(teacher_analytics, student_id)


@app.get("/api/teacher/course-lines/{course_line_id}/analytics")
async def course_line_analytics(course_line_id: str) -> dict:
    analytics_data = await asyncio.to_thread(teacher_analytics)
    analytics_data["courseLine"] = (await asyncio.to_thread(get_course_line, course_line_id)).model_dump()
    return analytics_data


# ---------- 教师 AI 分析 ----------

from backend.app.services import teacher_ai

@app.post("/api/teacher/ai/analyze-student")
async def analyze_student_endpoint(payload: dict) -> dict:
    student_id = payload.get("studentId")
    if not student_id:
        raise HTTPException(status_code=400, detail="studentId is required")
    try:
        return await teacher_ai.analyze_student_performance(student_id)
    except Exception as exc:
        return {"strengths": [], "weaknesses": ["分析失败"], "improvementPlan": ["请检查 AI 配置"], "recommendedTasks": [], "summary": str(exc)[:200]}


@app.post("/api/teacher/ai/analyze-class")
async def analyze_class_endpoint(payload: dict) -> dict:
    class_id = payload.get("classId")
    if not class_id:
        raise HTTPException(status_code=400, detail="classId is required")
    try:
        return await teacher_ai.analyze_class_performance(class_id)
    except Exception as exc:
        return {"commonWeaknesses": ["分析失败"], "topStudents": [], "strugglingStudents": [], "recommendations": [str(exc)[:200]]}


# ---------- 反馈系统 ----------

from backend.app.services.database import save_feedback, get_student_feedback, get_all_feedback

@app.post("/api/teacher/feedback/send")
async def send_feedback_endpoint(payload: dict) -> dict:
    teacher_id = payload.get("teacherId")
    student_id = payload.get("studentId")
    content = payload.get("content")
    analysis = payload.get("analysis")
    if not teacher_id or not student_id or not content:
        raise HTTPException(status_code=400, detail="teacherId, studentId and content are required")
    analysis_json = json.dumps(analysis, ensure_ascii=False) if analysis else None
    feedback_id = await asyncio.to_thread(save_feedback, teacher_id, student_id, content, analysis_json)
    return {"id": feedback_id, "ok": True}


@app.get("/api/student/feedback/{student_id}")
async def get_feedback_endpoint(student_id: str) -> dict:
    rows = await asyncio.to_thread(get_student_feedback, student_id)
    feedback = []
    for row in rows:
        item = dict(row)
        try:
            item["analysis"] = json.loads(item.pop("analysis_json", "null") or "null")
        except json.JSONDecodeError:
            item["analysis"] = None
        feedback.append(item)
    return {"feedback": feedback}


@app.get("/api/teacher/feedback/all")
async def get_all_feedback_endpoint() -> dict:
    rows = await asyncio.to_thread(get_all_feedback)
    feedback = []
    for row in rows:
        item = dict(row)
        try:
            item["analysis"] = json.loads(item.pop("analysis_json", "null") or "null")
        except json.JSONDecodeError:
            item["analysis"] = None
        feedback.append(item)
    return {"feedback": feedback}


# ---------- 错题本 API ----------

from backend.app.services.database import (
    add_error_entry,
    get_error_book,
    get_error_entry,
    update_error_entry,
    delete_error_entry,
    get_error_stats,
)

@app.get("/api/error-book/{student_id}")
async def get_error_book_endpoint(student_id: str, status: str | None = None) -> dict:
    entries = await asyncio.to_thread(get_error_book, student_id, status)
    return {"entries": entries, "total": len(entries)}


@app.get("/api/error-book/{student_id}/stats")
async def get_error_stats_endpoint(student_id: str) -> dict:
    return await asyncio.to_thread(get_error_stats, student_id)


@app.post("/api/error-book")
async def add_error_entry_endpoint(payload: dict) -> dict:
    entry_id = await asyncio.to_thread(
        add_error_entry,
        student_id=payload.get("studentId", ""),
        source_type=payload.get("sourceType", "manual"),
        question=payload.get("question", ""),
        student_answer=payload.get("studentAnswer", ""),
        correct_answer=payload.get("correctAnswer", ""),
        error_analysis=payload.get("errorAnalysis", ""),
        knowledge_points=payload.get("knowledgePoints", ""),
        difficulty=payload.get("difficulty", "medium"),
        source_id=payload.get("sourceId"),
        course_line_id=payload.get("courseLineId"),
        task_id=payload.get("taskId"),
    )
    return {"id": entry_id, "ok": True}


@app.put("/api/error-book/{entry_id}")
async def update_error_entry_endpoint(entry_id: int, payload: dict) -> dict:
    fields = {}
    if "status" in payload:
        fields["status"] = payload["status"]
    if "ai_analysis" in payload:
        fields["ai_analysis"] = payload["ai_analysis"]
    if "variant_question" in payload:
        fields["variant_question"] = payload["variant_question"]
    if "variant_answer" in payload:
        fields["variant_answer"] = payload["variant_answer"]
    ok = await asyncio.to_thread(update_error_entry, entry_id, **fields)
    return {"ok": ok}


@app.delete("/api/error-book/{entry_id}")
async def delete_error_entry_endpoint(entry_id: int) -> dict:
    ok = await asyncio.to_thread(delete_error_entry, entry_id)
    return {"ok": ok}


@app.post("/api/error-book/{entry_id}/analyze")
async def analyze_error_endpoint(entry_id: int) -> dict:
    """AI 分析错题，生成错误原因、知识点、正确思路。"""
    entry = await asyncio.to_thread(get_error_entry, entry_id)
    if not entry:
        raise HTTPException(status_code=404, detail="entry not found")

    prompt = f"""你是一个教学分析助手。请分析以下学生做错的题目。

题目：{entry['question']}
学生答案：{entry['student_answer'] or '未提供'}
正确答案：{entry['correct_answer'] or '未提供'}

请用 JSON 格式返回分析：
{{
  "errorAnalysis": "错误原因分析（具体说明学生哪里理解错了）",
  "knowledgePoints": "涉及的知识点（逗号分隔）",
  "correctApproach": "正确的解题思路",
  "tips": "避免再犯的建议"
}}"""

    try:
        result = await deepseek_client.direct_answer(prompt)
        # 提取 JSON 部分（兼容 markdown 代码块包裹）
        text = result.strip()
        if text.startswith("```"):
            text = text.split("\n", 1)[-1].rsplit("```", 1)[0].strip()
        analysis = json.loads(text)
        await asyncio.to_thread(
            update_error_entry, entry_id,
            ai_analysis=json.dumps(analysis, ensure_ascii=False),
            error_analysis=analysis.get("errorAnalysis", ""),
            knowledge_points=analysis.get("knowledgePoints", ""),
        )
        return {"ok": True, "analysis": analysis}
    except Exception as exc:
        return {"ok": False, "error": str(exc)[:200]}


@app.post("/api/error-book/{entry_id}/generate-variant")
async def generate_variant_endpoint(entry_id: int) -> dict:
    """基于错题生成变体练习题。"""
    entry = await asyncio.to_thread(get_error_entry, entry_id)
    if not entry:
        raise HTTPException(status_code=404, detail="entry not found")

    prompt = f"""你是一个出题助手。基于以下错题，生成一道类似的变体练习题。

原题：{entry['question']}
知识点：{entry['knowledge_points'] or '未标注'}

请用 JSON 格式返回：
{{
  "variantQuestion": "变体题目（和原题类似但不完全相同）",
  "variantAnswer": "参考答案"
}}"""

    try:
        result = await deepseek_client.direct_answer(prompt)
        text = result.strip()
        if text.startswith("```"):
            text = text.split("\n", 1)[-1].rsplit("```", 1)[0].strip()
        variant = json.loads(text)
        await asyncio.to_thread(
            update_error_entry, entry_id,
            variant_question=variant.get("variantQuestion", ""),
            variant_answer=variant.get("variantAnswer", ""),
        )
        return {"ok": True, "variant": variant}
    except Exception as exc:
        return {"ok": False, "error": str(exc)[:200]}


# ---------- 知识库文件管理 ----------

@app.get("/api/teacher/kb/files")
async def kb_files(course_line_id: str | None = None) -> dict:
    files = await asyncio.to_thread(kb_manager.list_files, course_line_id)
    return {"files": files, "total": len(files)}


@app.delete("/api/teacher/kb/files/{file_id}")
async def kb_delete_file(file_id: str) -> dict:
    deleted = await asyncio.to_thread(kb_manager.delete_file, file_id)
    if not deleted:
        raise HTTPException(status_code=404, detail="file not found")
    await asyncio.to_thread(rag_service.remove_document, file_id)
    return {"deleted": file_id}


@app.post("/api/teacher/kb/associate")
async def kb_associate(payload: dict) -> dict:
    file_id = payload.get("file_id")
    task_id = payload.get("task_id")
    if not file_id or not task_id:
        raise HTTPException(status_code=400, detail="file_id and task_id required")
    ok = await asyncio.to_thread(kb_manager.associate_task, file_id, task_id)
    task_ids = await asyncio.to_thread(kb_manager.get_file_tasks, file_id)
    await asyncio.to_thread(rag_service.set_document_tasks, file_id, task_ids)
    return {"ok": ok}


@app.delete("/api/teacher/kb/associate")
async def kb_dissociate(payload: dict) -> dict:
    file_id = payload.get("file_id")
    task_id = payload.get("task_id")
    if not file_id or not task_id:
        raise HTTPException(status_code=400, detail="file_id and task_id required")
    ok = await asyncio.to_thread(kb_manager.dissociate_task, file_id, task_id)
    task_ids = await asyncio.to_thread(kb_manager.get_file_tasks, file_id)
    await asyncio.to_thread(rag_service.set_document_tasks, file_id, task_ids)
    return {"ok": ok}


@app.get("/api/teacher/kb/task-files/{task_id}")
async def kb_task_files(task_id: str) -> dict:
    files = await asyncio.to_thread(kb_manager.get_task_files, task_id)
    return {"task_id": task_id, "files": files}


# ---------- 发布题目 ----------

from backend.app.services.database import (
    get_all_published_for_question,
    publish_question,
    unpublish_question,
    get_published_questions,
    get_published_questions_by_course,
)

@app.post("/api/teacher/questions/publish")
async def publish_question_endpoint(payload: dict) -> dict:
    question_id = payload.get("questionId")
    task_id = payload.get("taskId")
    course_line_id = payload.get("courseLineId")
    published_by = payload.get("publishedBy")
    if not question_id or not task_id or not course_line_id or not published_by:
        raise HTTPException(status_code=400, detail="missing required fields")
    row_id = await asyncio.to_thread(publish_question, question_id, task_id, course_line_id, published_by)
    return {"id": row_id, "ok": True}


@app.delete("/api/teacher/questions/unpublish")
async def unpublish_question_endpoint(payload: dict) -> dict:
    question_id = payload.get("questionId")
    task_id = payload.get("taskId")
    if not question_id or not task_id:
        raise HTTPException(status_code=400, detail="questionId and taskId required")
    ok = await asyncio.to_thread(unpublish_question, question_id, task_id)
    return {"ok": ok}


@app.get("/api/published-questions/{task_id}")
async def get_published_questions_endpoint(task_id: str) -> dict:
    rows = await asyncio.to_thread(get_published_questions, task_id)
    all_questions = _question_lookup()
    result = []
    for row in rows:
        q = all_questions.get(row["question_id"])
        if q:
            result.append(q.model_dump())
    return {"questions": result, "total": len(result)}


@app.get("/api/published-questions/course/{course_line_id}")
async def get_published_by_course_endpoint(course_line_id: str) -> dict:
    rows = await asyncio.to_thread(get_published_questions_by_course, course_line_id)
    all_questions = _question_lookup()
    tasks: dict[str, list] = {}
    for row in rows:
        task_id = row["task_id"]
        q = all_questions.get(row["question_id"])
        if q:
            tasks.setdefault(task_id, []).append(q.model_dump())
    return {"tasks": [{"taskId": tid, "questions": qs} for tid, qs in tasks.items()]}


@app.get("/api/questions/{question_id}/published-tasks")
async def get_question_published_tasks_endpoint(question_id: str) -> dict:
    rows = await asyncio.to_thread(get_all_published_for_question, question_id)
    return {
        "tasks": [
            {"task_id": row["task_id"], "course_line_id": row["course_line_id"]}
            for row in rows
        ]
    }


def _extract_pdf_text(target: Path) -> str:
    from pypdf import PdfReader
    reader = PdfReader(str(target))
    return "\n".join(page.extract_text() or "" for page in reader.pages)


def _question_from_teacher_row(row) -> Question:
    item = dict(row)
    options = None
    if item.get("options_json"):
        try:
            options = json.loads(item["options_json"])
        except json.JSONDecodeError:
            options = None
    try:
        tags = json.loads(item.get("tags_json") or "[]")
    except json.JSONDecodeError:
        tags = []
    return Question(
        id=item["id"],
        course_id=item["course_id"],
        lesson_id=item["lesson_id"],
        type=item["type"],
        stem=item["stem"],
        options=options,
        answer=item["answer"],
        explanation=item.get("explanation"),
        difficulty=item.get("difficulty") or "medium",
        tags=tags,
    )


def _extract_docx_text(target: Path) -> str:
    from docx import Document
    doc = Document(str(target))
    lines = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            lines.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                lines.append(" | ".join(cells))
    return "\n".join(lines)


def _extract_excel_text(target: Path) -> str:
    from openpyxl import load_workbook
    wb = load_workbook(str(target), read_only=True, data_only=True)
    lines = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        lines.append(f"=== Sheet: {sheet_name} ===")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            if any(cells):
                lines.append(" | ".join(cells))
    wb.close()
    return "\n".join(lines)


def _extract_csv_text(target: Path) -> str:
    import csv
    lines = []
    with open(str(target), "r", encoding="utf-8", errors="ignore") as f:
        reader = csv.reader(f)
        for row in reader:
            if any(cell.strip() for cell in row):
                lines.append(" | ".join(row))
    return "\n".join(lines)


def _sse(event: str, data: dict) -> str:
    return f"event: {event}\ndata: {json.dumps(data, ensure_ascii=False)}\n\n"


def _sse_response(event_stream) -> StreamingResponse:
    return StreamingResponse(
        event_stream,
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        },
    )


def _stream_metadata(state: dict) -> dict:
    steps = state.get("steps", [])
    current = int(state.get("current_step", 0))
    citations = state.get("citations", [])
    return {
        "sessionId": state.get("session_id", ""),
        "intent": state.get("intent", ""),
        "steps": steps,
        "currentStep": current,
        "totalSteps": len(steps),
        "currentStepTitle": steps[current]["title"] if steps and current < len(steps) else None,
        "citations": [citation.model_dump() for citation in citations],
        "status": state.get("status", "teaching"),
    }


def _learning_record_from_row(row) -> LearningRecord:
    item = dict(row)
    try:
        item["answers"] = json.loads(item.pop("answers_json") or "[]")
    except json.JSONDecodeError:
        item["answers"] = []
    return LearningRecord(**item)


def _safe_int(value, fallback: int = 0) -> int:
    try:
        return int(value)
    except (TypeError, ValueError):
        return fallback
